#!/usr/bin/env python3
"""Generate country labor law compliance guide Word document.
Uses raw NotebookLM answers — no DeepSeek polishing.
Usage: python3 build_word.py --answers answers.json --country "德国" --output ~/Desktop/德国劳动合规指南.docx
"""

import json, re, os, sys, argparse
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement
from docx.opc.constants import RELATIONSHIP_TYPE as RT

# ============================================================
# CLI
# ============================================================
parser = argparse.ArgumentParser()
parser.add_argument("--answers", required=True, help="Path to NotebookLM answers JSON")
parser.add_argument("--country", required=True, help="Country name in Chinese (e.g. 德国)")
parser.add_argument("--short", default="", help="Short country name for chapter titles")
parser.add_argument("--country-en", default="", help="Country name in English")
parser.add_argument("--summaries", default="", help="Path to chapter summaries JSON")
parser.add_argument("--firm", default="浩天律师事务所", help="Law firm name on cover")
parser.add_argument("--date", default="2026年5月", help="Cover date")
parser.add_argument("--output", required=True, help="Output .docx path")
parser.add_argument("--config", default="", help="YAML config for law table, preface, etc.")
args = parser.parse_args()

country = args.country
short = args.short or country
country_en = args.country_en or country

# Load answers
with open(args.answers) as f:
    ANSWERS = json.load(f)

# Load summaries
SUMMARIES = {}
if args.summaries and os.path.exists(args.summaries):
    with open(args.summaries) as f:
        SUMMARIES = json.load(f)

# Load config (optional YAML)
config = {}
if args.config and os.path.exists(args.config):
    import yaml
    with open(args.config) as f:
        config = yaml.safe_load(f)

doc = Document()

# ============================================================
# Page setup
# ============================================================
style = doc.styles['Normal']
style.font.name = '等线'
style.font.size = Pt(11)
style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
style.element.rPr.rFonts.set(qn('w:eastAsia'), '等线')

for section in doc.sections:
    section.page_width  = Mm(210)
    section.page_height = Mm(297)
    section.left_margin   = Inches(1.25)
    section.right_margin  = Inches(1.25)
    section.top_margin    = Inches(1.0)
    section.bottom_margin = Inches(1.0)

# ============================================================
# Helper functions
# ============================================================
def add_cover_title(doc, text, font_name='华文隶书', size=Pt(36), bold=False, color=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = font_name
    run.element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = size
    run.bold = bold
    if color: run.font.color.rgb = color
    return p

def add_section_header(doc, text, size=Pt(22), color=RGBColor(0x12, 0x41, 0x66)):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(24)
    run = p.add_run(text)
    run.font.name = '黑体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    run.font.size = size
    run.font.color.rgb = color
    return p

def add_sub_header(doc, text, size=Pt(16), bold=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = '黑体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    run.font.size = size
    run.bold = bold
    return p

def add_part_header(doc, text):
    return add_section_header(doc, text, Pt(22), RGBColor(0x12, 0x41, 0x66))

def add_chapter_header(doc, text):
    return add_sub_header(doc, text, Pt(16))

def add_body(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = '等线'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '等线')
    run.font.size = Pt(11)
    return p

def add_sub_heading(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = '黑体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    run.font.size = Pt(11)
    return p

def add_tip(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F2F2F2" w:val="clear"/>')
    pPr.append(shd)
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:left w:val="single" w:sz="12" w:space="4" w:color="124166"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)
    run = p.add_run(text)
    run.font.name = '等线'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '等线')
    run.font.size = Pt(11)
    return p

def add_note_header(doc, text="本章小结"):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = '黑体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    run.font.size = Pt(11)
    return p

def add_hyperlink_to_cell(cell, text, url, font_name='等线', font_size=Pt(11)):
    """Add clickable hyperlink in a table cell — the law name text IS the link."""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for r in p.runs:
        r.clear()
    part = p.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), font_name)
    rPr.append(rFonts)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(int(font_size.pt * 2)))
    rPr.append(sz)
    c = OxmlElement('w:color')
    c.set(qn('w:val'), '0563C1')
    rPr.append(c)
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    new_run.append(rPr)
    lines = text.split('\n')
    for i, line in enumerate(lines):
        if i > 0:
            new_run.append(OxmlElement('w:br'))
        t = OxmlElement('w:t')
        t.set(qn('xml:space'), 'preserve')
        t.text = line
        new_run.append(t)
    hyperlink.append(new_run)
    p._p.append(hyperlink)

# ============================================================
# Content processing
# ============================================================

def format_nlm_answer(text):
    """Format NotebookLM answer into paragraphs: subhead + body."""
    lines = text.split('\n')
    paragraphs = []
    current = ""
    for line in lines:
        stripped = line.strip()
        if not stripped:
            if current:
                paragraphs.append(('body', current.strip()))
                current = ""
            continue
        if re.match(r'^[一二三四五六七八九十]、', stripped) or re.match(r'^（[一二三四五六七八九十]）', stripped):
            if current:
                paragraphs.append(('body', current.strip()))
                current = ""
            paragraphs.append(('subhead', stripped))
        elif '内容：' in stripped or '合规建议：' in stripped:
            if current:
                paragraphs.append(('body', current.strip()))
                current = ""
            paragraphs.append(('body', stripped))
        else:
            current = current + stripped if current else stripped
    if current:
        paragraphs.append(('body', current.strip()))
    return paragraphs

def format_natural_text(text):
    """Convert NLM Q&A format into flowing natural paragraphs (for front matter)."""
    text = re.sub(r'\n(\d+)\n', r'[^\1]', text)
    paras = []
    current = ""
    for line in text.split('\n'):
        stripped = line.strip()
        if not stripped:
            if current:
                paras.append(current.strip())
                current = ""
            continue
        if re.match(r'^\d+$', stripped) and len(stripped) <= 3:
            continue
        cleaned = re.sub(r'^(一、|二、|三、|四、|五、|六、)\s*', '', stripped)
        cleaned = re.sub(r'^(内容：|合规建议：)', '', cleaned)
        current = current + cleaned if current else cleaned
    if current:
        paras.append(current.strip())
    return paras

def add_answer_section(doc, answer_text):
    for ptype, text in format_nlm_answer(answer_text):
        if ptype == 'subhead':
            add_sub_heading(doc, text)
        else:
            add_body(doc, text)

# ============================================================
# Law table (to be filled by caller; defaults to Canada example)
# ============================================================
def build_law_table(doc, law_table_data, law_font='等线'):
    """Build a 2-column law table with clickable hyperlinks in law name cells.
    law_table_data: list of (name, url, description) — url="" if no link.
    law_font: font for law name column (use 'Malgun Gothic' for Korean)."""
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for cell in table.columns[0].cells:
        cell.width = Cm(5.5)
    for cell in table.columns[1].cells:
        cell.width = Cm(11)
    hdr_cells = table.rows[0].cells
    for i, text in enumerate(['法律名称', '主要内容']):
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.name = '黑体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        run.font.size = Pt(11)
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="124166" w:val="clear"/>')
        hdr_cells[i]._tc.get_or_add_tcPr().append(shading)
    for law_name, law_url, law_desc in law_table_data:
        row = table.add_row()
        if law_url:
            add_hyperlink_to_cell(row.cells[0], law_name, law_url, font_name=law_font)
        else:
            p0 = row.cells[0].paragraphs[0]
            p0.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run0 = p0.add_run(law_name)
            run0.font.name = law_font
            run0.element.rPr.rFonts.set(qn('w:eastAsia'), law_font)
            run0.font.size = Pt(11)
        p1 = row.cells[1].paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run1 = p1.add_run(law_desc)
        run1.font.name = '等线'
        run1.element.rPr.rFonts.set(qn('w:eastAsia'), '等线')
        run1.font.size = Pt(11)
    doc.add_paragraph()

# ============================================================
# BUILD DOCUMENT
# ============================================================

# --- COVER ---
for _ in range(6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)

add_cover_title(doc, country, '华文隶书', Pt(36))
p = add_cover_title(doc, '劳动合规法律指南', '华文隶书', Pt(36))
p.paragraph_format.space_after = Pt(30)

add_cover_title(doc, 'Labor Law Compliance Guide', 'Times New Roman', Pt(22), bold=True)
add_cover_title(doc, f'for Chinese Investors in {country_en}', 'Times New Roman', Pt(22), bold=True)

for _ in range(4):
    doc.add_paragraph()

add_cover_title(doc, args.firm, '华文隶书', Pt(24))
p = add_cover_title(doc, args.date, '华文隶书', Pt(22))
p.paragraph_format.space_after = Pt(18)
doc.add_page_break()

# --- PREFACE (3 paragraphs; override from config if available) ---
add_section_header(doc, '前言')
prefaces = config.get('preface', {}).get('paragraphs', [
    f"随着中{short}经贸关系的持续深化，{country}凭借其优越的投资环境已成为中国企业海外布局的重要节点。然而，{country}的劳动法律体系具有独特性，对出海企业的用工管理提出了多维度合规要求。",
    f"本指南由{args.firm}跨境团队基于对{country}劳动法律体系的深入研究，并结合为中资企业提供{country}用工合规咨询的实务经验编纂而成。本指南从中国出海企业视角出发，全面梳理了{country}劳动法律框架、用工模式选择、劳动合同全生命周期管理、日常用工核心要素合规、专项合规管理、企业民主管理及劳动争议解决机制等关键议题。",
    f"需要特别说明的是，{country}劳动法律体系受联邦与地方多层级立法影响，具体合规方案须结合企业自身的业务模式、所在地区、用工规模及行业属性进行个案设计。建议读者在参考本指南的基础上，就具体法律问题咨询具备{country}执业资质的专业律师。",
])
for para in prefaces:
    add_body(doc, para)

# --- FDI Environment ---
add_sub_header(doc, f'{country}外商投资环境')
for para in format_natural_text(ANSWERS.get("外商投资环境", "")):
    if para:
        add_body(doc, para)

# --- Labor Compliance Overview ---
add_sub_header(doc, f'{country}劳动合规概要')
for para in format_natural_text(ANSWERS.get("劳动合规概要", "")):
    if para:
        add_body(doc, para)

# --- Legal System Table ---
add_sub_header(doc, f'{country}劳动法律体系')
law_table_data = config.get('legal_system', {}).get('laws', [])
if law_table_data:
    law_font = config.get('legal_system', {}).get('font', '等线')
    build_law_table(doc, law_table_data, law_font=law_font)
else:
    add_body(doc, f"（{country}劳动法律体系表待配置。请在 config 中提供 legal_system.laws 列表。）")
    doc.add_paragraph()

# --- TOC ---
doc.add_page_break()
add_section_header(doc, '目录')
toc_entries = [
    (1, '第一部分 用工主体与用工模式合规'),
    (2, f'第1章 外商投资企业在{short}用工主体资格合规'),
    (2, '第2章 用工形式选择'),
    (1, '第二部分 劳动合同全生命周期管理'),
    (2, '第3章 劳动合同订立合规'),
    (2, '第4章 劳动合同履行与变更合规'),
    (2, '第5章 劳动合同解除与终止合规'),
    (1, '第三部分 日常用工核心要素合规'),
    (2, '第6章 工作时间与休息休假合规'),
    (2, '第7章 工作薪酬合规管理'),
    (2, '第8章 公积金及保险合规'),
    (2, '第9章 劳动保护与职业安全合规'),
    (1, '第四部分 专项合规管理'),
    (2, '第10章 商业秘密保护与竞业限制合规'),
    (2, '第11章 特殊员工保护合规'),
    (2, '第12章 反歧视与反性骚扰合规'),
    (1, '第五部分 企业民主管理与劳资协同合规'),
    (2, f'第13章 {short}工会运作合规'),
    (2, '第14章 企业内部民主管理合规'),
    (1, '第六部分 劳动争议解决全流程应对'),
    (2, f'第15章 {short}劳动争议解决机制'),
]
for level, entry in toc_entries:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(entry)
    run.font.name = '等线'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '等线')
    run.font.size = Pt(11) if level == 1 else Pt(10.5)
    if level == 1:
        run.bold = True

# ============================================================
# CHAPTER CONTENT
# ============================================================
parts = [
    ('第一部分 用工主体与用工模式合规', [
        (f'第1章 外商投资企业在{short}用工主体资格合规', '外商投资企业用工主体资格合规'),
        ('第2章 用工形式选择', '用工形式选择'),
    ]),
    ('第二部分 劳动合同全生命周期管理', [
        ('第3章 劳动合同订立合规', '劳动合同订立合规'),
        ('第4章 劳动合同履行与变更合规', '劳动合同履行与变更合规'),
        ('第5章 劳动合同解除与终止合规', '劳动合同解除与终止合规'),
    ]),
    ('第三部分 日常用工核心要素合规', [
        ('第6章 工作时间与休息休假合规', '工作时间与休息休假合规'),
        ('第7章 工作薪酬合规管理', '工作薪酬合规管理'),
        ('第8章 公积金及保险合规', '公积金及保险合规'),
        ('第9章 劳动保护与职业安全合规', '劳动保护与职业安全合规'),
    ]),
    ('第四部分 专项合规管理', [
        ('第10章 商业秘密保护与竞业限制合规', '商业秘密保护与竞业限制合规'),
        ('第11章 特殊员工保护合规', '特殊员工保护合规'),
        ('第12章 反歧视与反性骚扰合规', '反歧视与反性骚扰合规'),
    ]),
    ('第五部分 企业民主管理与劳资协同合规', [
        (f'第13章 {short}工会运作合规', '工会运作合规'),
        ('第14章 企业内部民主管理合规', '企业内部民主管理合规'),
    ]),
    ('第六部分 劳动争议解决全流程应对', [
        (f'第15章 {short}劳动争议解决机制', '劳动争议解决机制'),
    ]),
]

for part_name, chapters in parts:
    doc.add_page_break()
    add_part_header(doc, part_name)

    for idx, (chapter_title, answer_key) in enumerate(chapters):
        if idx > 0:
            doc.add_page_break()
        add_chapter_header(doc, chapter_title)
        answer = ANSWERS.get(answer_key, '')
        if not answer:
            for k, v in ANSWERS.items():
                if answer_key in k or k in answer_key:
                    answer = v
                    break

        if answer:
            add_answer_section(doc, answer)

            # Chapter summary
            add_note_header(doc, '本章小结')
            summary_key = f"CA:{chapter_title}"
            summary_text = SUMMARIES.get(summary_key, "")
            if summary_text:
                add_body(doc, summary_text)
            else:
                add_body(doc, f"（本章小结待生成——请运行 generate_summaries.py）")
        else:
            add_body(doc, f"（{answer_key}相关内容待补充）")

# --- DISCLAIMER / FIRM INFO ---
doc.add_page_break()

def add_end_section_header(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(16)
    run.bold = True
    return p

def add_end_body(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)
    return p

disclaimer_default = (
    '浩天律师事务所编写《' + country + '劳动合规法律指南》'
    '（以下简称“指南”），仅为帮助中企及时了解' + country + '劳动领域法律、政策及实务最新动态，'
    '助力企业防范用工合规风险、规范用工管理。“指南”所载信息基于现行劳动法律法规、公开政策文件及行业实践整理，'
    '仅供参考交流，不构成浩天律师事务所及其律师对相关问题的正式法律意见。'
)
disclaimer_paras = config.get('disclaimer', {}).get('paragraphs', [
    disclaimer_default,
    '鉴于' + country + '劳动法律法规可能发生修订、政策文件可能调整、司法实践存在差异，'
    '且具体用工情境下的法律适用受多种因素影响，读者不应仅依据“指南”内容做出商业决策或人事处理决定。'
    '如需针对具体用工事项寻求法律意见，建议咨询具备' + country + '执业资质的专业律师或当地官方机构。',
    '浩天律师事务所不为“指南”非经授权转载、引述或修改所导致的误读、遗漏或偏差承担责任。'
    '“指南”著作权归浩天律师事务所所有，未经书面许可，任何单位或个人不得以营利为目的复制、传播或修改。',
])

add_end_section_header(doc, '特别声明')
for para in disclaimer_paras:
    add_end_body(doc, para)

doc.add_paragraph()

contact_info = config.get('contact', {})
contact_name = contact_info.get('name', '梁马玲')
contact_title = contact_info.get('title', '高级合伙人｜浩天全国劳动法专业委员会牵头合伙人')
contact_office = contact_info.get('office', '办公室：浩天律师事务所｜上海总部')
contact_email = contact_info.get('email', 'Email:   marinda.liang@hylandslaw.com')

add_end_section_header(doc, '联系人')
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(6)
p.paragraph_format.space_after = Pt(6)
run = p.add_run(contact_name + '    ' + contact_title)
run.font.name = 'Times New Roman'
run.font.size = Pt(11)
run.bold = True

add_end_body(doc, contact_office)
add_end_body(doc, contact_email)

doc.add_paragraph()

firm_intro_paras = config.get('firm_intro', {}).get('paragraphs', [
    '北京浩天律师事务所（“浩天”）创立于1997年，双总部分别位于北京和上海。'
    '历经多年发展，浩天现已设有北京、上海、广州、深圳、长沙、成都、重庆、大连、福州、贵阳、哈尔滨、海口、杭州、'
    '合肥、呼和浩特、济南、昆明、南昌、南京、南宁、宁波、青岛、石家庄、沈阳、苏州、天津、太原、温州、武汉、'
    '乌鲁木齐、西安、银川、郑州、香港等34家办公室，并在全球130多个国家和地区拥有150余家合作律师事务所。',
    '浩天现有600余位合伙人及顾问，2000余名律师和专业人员，能够为客户提供英语、日语、韩语、德语、法语、'
    '意大利语等多种语言的法律服务。浩天是中国优秀的大型综合性高端商事律师事务所之一，'
    '连续多年荣获钱伯斯(Chambers & Partners)、亚洲法律杂志(ALB)、The Legal 500、LegalBand等国际知名法律服务评级机构的推荐及奖项。',
    '浩天各业务团队均具备扎实的专业技能和丰富的项目经验，他们不仅是深谙各个领域的法律专家，'
    '而且熟悉中国的投资环境和各个行业的不同商业惯例，拥有深厚的实务经验。'
    '浩天通过紧密高效的团队合作，能够针对客户的不同需求，提供量身定制的法律解决方案。',
])

add_end_section_header(doc, '浩天简介')
for para in firm_intro_paras:
    add_end_body(doc, para)

# --- SAVE ---
doc.save(args.output)
print(f"Document saved to: {args.output}")
print(f"File size: {os.path.getsize(args.output)} bytes")
