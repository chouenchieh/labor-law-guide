#!/usr/bin/env python3
"""Generate Chinese law reference Word doc from DeepSeek translations.
Usage: python3 build_law_reference.py --translations translations.json --font "Malgun Gothic" --output ref.docx
"""
import json, argparse
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

parser = argparse.ArgumentParser()
parser.add_argument("--translations", required=True)
parser.add_argument("--font", default="等线", help="Font for original law name text")
parser.add_argument("--title", default="法律法规中文对照表", help="Document title")
parser.add_argument("--output", required=True)
args = parser.parse_args()

with open(args.translations) as f:
    translations = json.load(f)

doc = Document()

# Page setup
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

style = doc.styles['Normal']
font = style.font
font.name = '等线'
font.size = Pt(11)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run(args.title)
run.font.size = Pt(16)
run.font.bold = True
run.font.name = '黑体'

doc.add_paragraph()

for i, text in enumerate(translations):
    lines = text.strip().split('\n')
    if not lines:
        continue

    # Law name (original language)
    p = doc.add_paragraph()
    run = p.add_run(lines[0].strip())
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.name = args.font

    for line in lines[1:]:
        line = line.strip()
        if not line:
            continue
        p = doc.add_paragraph()
        if line.startswith("中文名称：") or line.startswith("中文名称:"):
            run = p.add_run(line)
            run.font.size = Pt(11)
            run.font.name = '等线'
        elif line.startswith("简介：") or line.startswith("简介:"):
            run = p.add_run(line)
            run.font.size = Pt(10)
            run.font.name = '等线'
        elif line.startswith("链接：") or line.startswith("链接:"):
            run = p.add_run(line)
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x42, 0x6E, 0xB4)
            run.font.name = '等线'
        else:
            run = p.add_run(line)
            run.font.size = Pt(10)
            run.font.name = '等线'

    # Space between laws
    if i < len(translations) - 1:
        doc.add_paragraph()

doc.save(args.output)
print(f"Law reference saved to: {args.output}")
